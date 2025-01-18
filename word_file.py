from docx import Document

document=Document()
document.add_heading("Hello World")
p=document.add_paragraph("Hello this is a sample tetxt")
p.add_run("This text is bold").bold=True
p.add_run("This text is bold").italic=True

document.add_paragraph('This is item one',style="List Bullet")
document.add_paragraph('This is item two',style="List Bullet")
document.add_paragraph('This is item three',style="List Bullet")

table_header = ["Name","Age","Job"]

some_data = [
    ["John", 25, "Engineer"],
    ["Jane", 30, "Doctor"],
    ["Bob", 28, "Teacher"],
    ["Alice", 27, "Designer"]
    
]
table = document.add_table(rows=1,cols=3)
for i in range(3):
    table.rows[0].cells[i].text = table_header[i]

for name,age,job in some_data:
    cells = table.add_row().cells
    cells[0].text = name
    cells[1].text = str(age)
    cells[2].text = job
    


document.add_page_break()

document.add_paragraph("Wekcome to a new page")
document.add_picture("money.jpg")
document.save("test.docx")
